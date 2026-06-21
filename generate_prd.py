from docx import Document
from docx.shared import Pt

doc = Document()

# Title
section = doc.sections[0]
section.top_margin = Pt(36)
section.bottom_margin = Pt(36)
section.left_margin = Pt(36)
section.right_margin = Pt(36)

doc.add_heading('Product Requirements Document', level=0)
doc.add_paragraph('TukarSampah')
doc.add_paragraph('Aplikasi Bank Sampah & Tukar Poin Digital')
doc.add_paragraph('Tanggal: 2026-06-21')
doc.add_paragraph('')

# Overview
_doc = doc.add_heading('1. Ringkasan Proyek', level=1)
doc.add_paragraph(
    'TukarSampah adalah aplikasi mobile berbasis Flutter yang menghubungkan pengguna dengan layanan bank sampah digital, ' 
    'memungkinkan pengguna mengonversi sampah menjadi poin, menjadwalkan penjemputan, dan menukarkan poin dengan voucher digital.'
)

doc.add_heading('2. Tujuan', level=2)
doc.add_paragraph('Tujuan utama aplikasi adalah:')
for item in [
    'Mendorong pengelolaan sampah yang lebih baik dengan imbalan poin.',
    'Menyediakan pengalaman digital yang mudah untuk menjadwalkan penjemputan sampah.',
    'Memfasilitasi penukaran poin menjadi voucher pulsa, token listrik, dan saldo e-wallet.',
    'Memberikan panel admin untuk mengelola katalog, penjemputan, dan pengguna.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# Stakeholders

doc.add_heading('3. Pemangku Kepentingan', level=2)
doc.add_paragraph('Pemangku kepentingan utama:')
for item in [
    'Pengguna akhir (warga) yang ingin menukarkan sampah dengan poin dan rewards.',
    'Admin operasional yang mengelola item katalog, jadwal penjemputan, dan pengguna.',
    'Pengembang yang memelihara aplikasi Flutter dan backend Supabase.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# Scope

doc.add_heading('4. Ruang Lingkup', level=1)
doc.add_paragraph('Fitur yang termasuk dalam ruang lingkup:')
for item in [
    'Registrasi dan login pengguna dengan Supabase Authentication.',
    'Dashboard pengguna dengan ringkasan poin dan akses cepat ke fitur utama.',
    'Kalkulator poin berdasarkan kategori sampah dan berat.',
    'Penjadwalan penjemputan sampah dengan tanggal, waktu, alamat, jenis sampah, dan estimasi berat.',
    'Katalog penukaran poin dengan item aktif yang dapat ditukar.',
    'Halaman profil dengan detail pengguna dan riwayat transaksi.',
    'Panel admin untuk statistik, pengelolaan katalog, penjemputan, dan pengguna.',
    'Backend Supabase dengan RLS dan tabel profiles, pickup_schedules, catalog_items, transactions.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

doc.add_paragraph('Fitur di luar ruang lingkup saat ini:')
for item in [
    'Integrasi pembayaran otomatis untuk item katalog.',
    'Fitur chat atau notifikasi real-time push.',
    'Akun keluarga atau multi-user dengan satu alamat yang membagikan poin.',
    'Dashboard analitik advanced untuk pelaporan operasional.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# User Roles

doc.add_heading('5. Peran Pengguna', level=1)
doc.add_heading('5.1 Pengguna Biasa', level=2)
doc.add_paragraph('Kemampuan:')
for item in [
    'Register dan login ke aplikasi.',
    'Melihat total poin dan kategori sampah.',
    'Menghitung estimasi poin dari berat sampah.',
    'Menjadwalkan penjemputan sampah.',
    'Menukar poin melalui katalog item.',
    'Melihat profil dan riwayat transaksi.',
    'Logout dari akun.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

doc.add_heading('5.2 Admin', level=2)
doc.add_paragraph('Kemampuan:')
for item in [
    'Melihat statistik aplikasi: total user, total jemput, jadwal pending, dan total item katalog.',
    'Mengelola katalog item reward: tambah, edit, hapus.',
    'Mengelola jadwal penjemputan: ubah status penjemputan.',
    'Melihat data pengguna dan mengubah total poin pengguna.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# Functional requirements

doc.add_heading('6. Kebutuhan Fungsional', level=1)

# Authentication

doc.add_heading('6.1 Autentikasi', level=2)
doc.add_paragraph('Pengguna harus dapat melakukan registrasi dan login menggunakan email dan password melalui Supabase Authentication. ' 
                  'Setelah login, peran pengguna ditentukan dari tabel profiles untuk menavigasi ke halaman /home atau /admin.')
doc.add_paragraph('Aturan validasi:')
for item in [
    'Email harus berisi @ dan tidak boleh kosong.',
    'Password minimal 6 karakter.',
    'Nama lengkap, telepon, dan alamat wajib di registrasi.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# Home dashboard

doc.add_heading('6.2 Dashboard Pengguna', level=2)
doc.add_paragraph('Halaman beranda menampilkan greeting, total poin pengguna, aksi cepat, kategori sampah, dan informasi ringkas. ' 
                  'Pengguna dapat mengakses kalkulator poin, jadwal penjemputan, katalog, dan profil dari bottom navigation.')

# calculator

doc.add_heading('6.3 Kalkulator Poin', level=2)
doc.add_paragraph('Pengguna dapat memilih jenis sampah dan memasukkan berat untuk menghitung estimasi poin menggunakan konversi berikut:')
doc.add_paragraph('Plastik 50 poin/kg, Kertas 30 poin/kg, Elektronik 200 poin/kg, Logam 100 poin/kg, Kaca 40 poin/kg.')
doc.add_paragraph('Hasil menghitung ditampilkan pada kartu hasil dan rincian perhitungan disimpan sementara dalam sesi layar kalkulator.')

# pickup

doc.add_heading('6.4 Penjadwalan Penjemputan', level=2)
doc.add_paragraph('Pengguna dapat membuat jadwal penjemputan dengan memilih tanggal, waktu, alamat, jenis sampah, estimasi berat, dan catatan opsional. ' 
                  'Data disimpan ke tabel pickup_schedules dengan status awal pending.')
doc.add_paragraph('Halaman penjemputan menampilkan daftar jadwal pengguna sendiri, status, dan detail ringkas. Pengguna dapat merefresh daftar melalui swipe down.')

# catalog

doc.add_heading('6.5 Katalog Penukaran', level=2)
doc.add_paragraph('Pengguna dapat melihat item katalog aktif yang tersedia untuk ditukar. Item dibagi dalam tab Pulsa, Token Listrik, dan E-Wallet.')
doc.add_paragraph('Jika poin pengguna cukup, pengguna dapat menukar item. Penukaran akan mengurangi total poin pengguna pada tabel profiles dan mencatat transaksi dengan tipe redeem di tabel transactions.')

# profile

doc.add_heading('6.6 Profil & Riwayat', level=2)
doc.add_paragraph('Halaman profil menampilkan informasi pengguna, statistik total poin, total poin diperoleh, total poin ditukar, dan riwayat transaksi terbaru. ' 
                  'Pengguna dapat logout dari akun.')

# admin

doc.add_heading('6.7 Panel Admin', level=2)
doc.add_paragraph('Admin memiliki akses ke panel dashboard admin dengan navigasi ke:')
for item in [
    'Dashboard utama dengan statistik operasional.',
    'Kelola katalog item reward.',
    'Kelola jadwal penjemputan dan update status.',
    'Kelola daftar pengguna dan update total poin pengguna.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# Data model

doc.add_heading('7. Model Data', level=1)
doc.add_paragraph('Arsitektur backend menggunakan Supabase dengan tabel-tabel utama berikut:')

# Table: profiles
if True:
    doc.add_heading('7.1 profiles', level=2)
    doc.add_paragraph('Tabel menyimpan data pengguna dan poin.')
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kolom'
    hdr_cells[1].text = 'Tipe'
    hdr_cells[2].text = 'Deskripsi'
    rows = [
        ('id', 'UUID', 'Primary key, referensi auth.users(id)'),
        ('email', 'TEXT', 'Email pengguna'),
        ('nama_lengkap', 'TEXT', 'Nama lengkap'),
        ('no_telepon', 'TEXT', 'Nomor telepon'),
        ('alamat', 'TEXT', 'Alamat'),
        ('total_poin', 'INTEGER', 'Saldo poin'),
        ('created_at', 'TIMESTAMP', 'Waktu pembuatan record'),
        ('role', 'TEXT', 'Peran pengguna, digunakan oleh aplikasi (admin/user)'),
    ]
    for col, tipe, desc in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = tipe
        row_cells[2].text = desc

# Table: pickup_schedules
if True:
    doc.add_heading('7.2 pickup_schedules', level=2)
    doc.add_paragraph('Tabel menyimpan jadwal penjemputan yang dibuat oleh pengguna.')
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kolom'
    hdr_cells[1].text = 'Tipe'
    hdr_cells[2].text = 'Deskripsi'
    rows = [
        ('id', 'UUID', 'Primary key'),
        ('user_id', 'UUID', 'Referensi auth.users(id)'),
        ('tanggal_jemput', 'DATE', 'Tanggal penjemputan'),
        ('waktu_jemput', 'TEXT', 'Slot waktu penjemputan'),
        ('alamat_jemput', 'TEXT', 'Alamat penjemputan'),
        ('catatan', 'TEXT', 'Catatan opsional'),
        ('jenis_sampah', 'TEXT[]', 'Daftar jenis sampah'),
        ('estimasi_berat', 'DECIMAL', 'Estimasi berat dalam kg'),
        ('status', 'TEXT', 'Status: pending/confirmed/onTheWay/completed/cancelled'),
        ('created_at', 'TIMESTAMP', 'Waktu pembuatan record'),
    ]
    for col, tipe, desc in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = tipe
        row_cells[2].text = desc

# Table: catalog_items
if True:
    doc.add_heading('7.3 catalog_items', level=2)
    doc.add_paragraph('Tabel menyimpan item reward yang dapat ditukar pengguna.')
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kolom'
    hdr_cells[1].text = 'Tipe'
    hdr_cells[2].text = 'Deskripsi'
    rows = [
        ('id', 'UUID', 'Primary key'),
        ('nama', 'TEXT', 'Nama item reward'),
        ('deskripsi', 'TEXT', 'Deskripsi item'),
        ('poin_dibutuhkan', 'INTEGER', 'Poin yang harus dibayar'),
        ('kategori', 'TEXT', 'Kategori item: pulsa/tokenListrik/eWallet'),
        ('nominal', 'TEXT', 'Nominal atau nilai voucher'),
        ('stok', 'INTEGER', 'Jumlah stok item'),
        ('image_url', 'TEXT', 'URL gambar (opsional)'),
        ('is_active', 'BOOLEAN', 'Status aktif item'),
        ('created_at', 'TIMESTAMP', 'Waktu pembuatan record'),
    ]
    for col, tipe, desc in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = tipe
        row_cells[2].text = desc

# Table: transactions
if True:
    doc.add_heading('7.4 transactions', level=2)
    doc.add_paragraph('Tabel menyimpan riwayat transaksi poin pengguna.')
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kolom'
    hdr_cells[1].text = 'Tipe'
    hdr_cells[2].text = 'Deskripsi'
    rows = [
        ('id', 'UUID', 'Primary key'),
        ('user_id', 'UUID', 'Referensi auth.users(id)'),
        ('tipe', 'TEXT', 'Tipe transaksi: deposit atau redeem'),
        ('poin', 'INTEGER', 'Jumlah poin transaksi'),
        ('deskripsi', 'TEXT', 'Deskripsi transaksi'),
        ('created_at', 'TIMESTAMP', 'Waktu pembuatan record'),
    ]
    for col, tipe, desc in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = tipe
        row_cells[2].text = desc

# Tech stack

doc.add_heading('8. Teknologi', level=1)
doc.add_paragraph('Teknologi dan dependensi yang digunakan:')
for item in [
    'Flutter (Dart) untuk aplikasi mobile cross-platform.',
    'Riverpod untuk state management.',
    'GoRouter untuk navigasi dan route guard.',
    'Supabase untuk backend: authentication, database, dan row-level security.',
    'Google Fonts Poppins untuk tampilan teks modern.',
    'intl untuk format tanggal lokal.',
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

doc.add_heading('9. Arsitektur Navigasi', level=1)
doc.add_paragraph('Struktur navigasi aplikasi:')
for item in [
    '/login, /register untuk autentikasi.',
    '/home, /calculator, /pickup, /catalog, /profile untuk pengguna biasa.',
    '/schedule-pickup untuk membuat jadwal penjemputan.',
    '/admin, /admin/catalog, /admin/pickups, /admin/users untuk admin.',
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# Non-functional requirements

doc.add_heading('10. Kebutuhan Non-Fungsional', level=1)
for item in [
    'Aplikasi harus responsif pada perangkat Android dan iOS.',
    'Navigasi harus cepat dan konsisten menggunakan bottom navigation dan shell routes.',
    'Data sensitif disimpan dan diakses melalui Supabase dengan kontrol RLS.',
    'Validasi input wajib dilakukan sebelum permintaan ke backend.',
    'Antarmuka menggunakan Material 3 dan skema warna hijau untuk tema ramah lingkungan.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# Risks and dependencies

doc.add_heading('11. Risiko dan Dependensi', level=1)
for item in [
    'Ketersediaan Supabase dan koneksi internet akan memengaruhi fungsi utama aplikasi.',
    'Penggunaan kunci anon key di klien membutuhkan izin database yang tepat di Supabase.',
    'Penanganan status penjemputan memerlukan sinkronisasi antara admin dan pengguna.',
    'Pengembangan lebih lanjut harus menambahkan notifikasi dan manajemen inventaris yang lebih lengkap.'
]:
    doc.add_paragraph(f'- {item}', style='List Bullet')

# Appendix

doc.add_heading('12. Lampiran', level=1)
doc.add_paragraph('Skema database dan konversi poin dapat diadaptasi sesuai kebutuhan operasional. ' 
                  'Implementasi saat ini mencakup sample data katalog untuk pulsa, token listrik, dan e-wallet.')

doc.save('TukarSampah_PRD.docx')
print('Generated TukarSampah_PRD.docx')
